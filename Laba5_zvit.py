"""
Скрипт для автоматичної генерації звіту з лабораторної роботи №5
Аспірант: Косарєв Валерій Геннадійович
Напрямок: 051 Економіка
Тема: Побудова рекомендаційної системи відеопереглядів
"""

import os
import subprocess
import sys
from pathlib import Path
from datetime import datetime
import ast
import importlib
from io import StringIO
import contextlib

# Перевірка наявності необхідних бібліотек
try:
	from docx import Document  # type: ignore
	from docx.shared import Inches, Pt, RGBColor  # type: ignore
	from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
except ImportError:
	print("Встановлення бібліотеки python-docx...")
	subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
	from docx import Document  # type: ignore
	from docx.shared import Inches, Pt, RGBColor  # type: ignore
	from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

# Конфігурація
CURRENT_DIR = Path(__file__).parent
DATA_FILE = CURRENT_DIR / "База даних для виконання лабораторної роботи.txt"

REQUIRED_PACKAGES = [
	"pandas",
	"numpy",
	"matplotlib",
	"seaborn",
	"scipy",
]

# Описи завдань
TASK_DESCRIPTIONS = {
	'zada4a1.py': {
		'title': 'Завдання 1: Завантаження та аналіз даних',
		'description': 'Завантаження даних про відеопереглядання, обчислення статистик та аналіз структури даних.',
		'topic': 'Завантаження даних, аналіз, статистика'
	},
	'zada4a2.py': {
		'title': 'Завдання 2: Нормалізація даних',
		'description': 'Застосування різних методів нормалізації даних для рекомендаційної системи.',
		'topic': 'Нормалізація, Z-score, Min-Max'
	},
	'zada4a3.py': {
		'title': 'Завдання 3: Обчислення метрик відстані',
		'description': 'Розрахунок евклідової, Манхеттенської та Max-метрики для порівняння користувачів.',
		'topic': 'Метрики, відстані, подібність'
	},
	'zada4a4.py': {
		'title': 'Завдання 4: Побудова рекомендацій',
		'description': 'Розробка рекомендаційної системи на основі метрик подібності користувачів.',
		'topic': 'Рекомендаційна система, KNN'
	},
	'zada4a5.py': {
		'title': 'Завдання 5: Формування результатів із ключовими словами',
		'description': 'Підготовка рекомендацій з додатковою інформацією та ключовими словами для презентації.',
		'topic': 'Обробка результатів, аналіз'
	},
}

TEST_INPUTS = {
    'zada4a1.py': 'G:\\My Drive\\Aspirantura\\ХАІ Економіка асп\\01. Сучасні економічні теорії\\Машинне навчання, обробка великих масивів даних\\Laba5\\База даних для виконання лабораторної роботи.txt\n1\nn\n',
	'zada4a2.py': 'G:\\My Drive\\Aspirantura\\ХАІ Економіка асп\\01. Сучасні економічні теорії\\Машинне навчання, обробка великих масивів даних\\Laba5\\База даних для виконання лабораторної роботи.txt\n1\nn\n',
	'zada4a3.py': 'G:\\My Drive\\Aspirantura\\ХАІ Економіка асп\\01. Сучасні економічні теорії\\Машинне навчання, обробка великих масивів даних\\Laba5\\База даних для виконання лабораторної роботи.txt\n1\nn\n',
	'zada4a4.py': 'G:\\My Drive\\Aspirantura\\ХАІ Економіка асп\\01. Сучасні економічні теорії\\Машинне навчання, обробка великих масивів даних\\Laba5\\База даних для виконання лабораторної роботи.txt\n1\nn\n',
	'zada4a5.py': 'G:\\My Drive\\Aspirantura\\ХАІ Економіка асп\\01. Сучасні економічні теорії\\Машинне навчання, обробка великих масивів даних\\Laba5\\База даних для виконання лабораторної роботи.txt\n1\nn\n',
}


def ensure_packages() -> None:
	"""Перевіряє та встановлює необхідні пакети."""
	missing = []
	for pkg in REQUIRED_PACKAGES:
		try:
			importlib.import_module(pkg)
		except ImportError:
			missing.append(pkg)

	if missing:
		print(f"Встановлення пакетів: {', '.join(missing)}")
		subprocess.check_call([sys.executable, "-m", "pip", "install", *missing])


def get_python_files() -> list:
	"""Збирає всі Python файли у папці, окрім звіту"""
	files = []
	for file in CURRENT_DIR.glob("*.py"):
		if file.name != "Laba5_zvit.py":
			files.append(file)
	return sorted(files, key=lambda x: x.name)


def capture_output(script_path, input_text=None, args=None):
	"""Виконує скрипт та захоплює його вивід"""
	try:
		env = os.environ.copy()
		env['PYTHONIOENCODING'] = 'utf-8'
		env['MPLBACKEND'] = 'Agg'

		cmd = [sys.executable, str(script_path)]
		if args:
			cmd.extend(args)

		# Якщо є вхідні дані, передаємо їх через stdin
		if input_text:
			result = subprocess.run(
				cmd,
				capture_output=True,
				timeout=60,
				input=input_text.encode('utf-8'),
				cwd=str(script_path.parent),
				env=env,
			)
			# Декодуємо вихід з заміною невалідних символів
			stdout_text = result.stdout.decode('utf-8', errors='replace').strip()
			stderr_text = result.stderr.decode('utf-8', errors='replace').strip()
		else:
			result = subprocess.run(
				cmd,
				capture_output=True,
				timeout=60,
				cwd=str(script_path.parent),
				env=env,
			)
			stdout_text = result.stdout.decode('utf-8', errors='replace').strip()
			stderr_text = result.stderr.decode('utf-8', errors='replace').strip()
		
		if result.returncode == 0:
			return stdout_text if stdout_text else stderr_text
		if stdout_text or stderr_text:
			combined = "\n".join([t for t in [stdout_text, stderr_text] if t])
			return combined
		return None
	except subprocess.TimeoutExpired:
		return None
	except Exception:
		return None


def read_file_content(file_path: Path) -> str:
	"""Читає вміст файлу"""
	try:
		return file_path.read_text(encoding='utf-8')
	except Exception as e:
		return f"Помилка читання файлу: {str(e)}"


def add_formatted_code(doc: Document, code_text: str) -> None:
	"""Додає форматований код до документа без скорочень"""
	for line in code_text.split('\n'):
		p = doc.add_paragraph(line if line.strip() else "")
		p.paragraph_format.left_indent = Inches(0.5)
		p.paragraph_format.line_spacing = 1.0
		if p.runs:
			p.runs[0].font.name = 'Consolas'
			p.runs[0].font.size = Pt(8)
			p.runs[0].font.color.rgb = RGBColor(0, 0, 0)


def get_task_conclusions(task_file: str) -> str:
	"""Повертає висновки для конкретного завдання"""
	conclusions = {
		'zada4a1.py': 'Виконано завантаження та первинний аналіз даних про відеопереглядання користувачів.',
		'zada4a2.py': 'Застосовано методи нормалізації даних для підготовки до розрахунку метрик.',
		'zada4a3.py': 'Розраховано метрики відстані для оцінки подібності користувачів та їхніх переглядань.',
		'zada4a4.py': 'Побудовано рекомендаційну систему на основі метрик подібності користувачів.',
		'zada4a5.py': 'Сформовано остаточні рекомендації з додатковою інформацією та ключовими словами.',
	}
	return conclusions.get(task_file, 'Завдання виконано успішно.')


def find_task_images(task_name: str) -> list[Path]:
	"""Знаходить зображення, пов'язані з конкретним завданням"""
	# Шукаємо всі PNG файли у папці
	all_images = sorted(CURRENT_DIR.glob('*.png'))
	
	if task_name == 'zada4a1.py':
		# Для першого завдання повертаємо перші кілька зображень, якщо вони є
		return all_images[:3] if all_images else []
	
	return []


def add_task_images_section(doc: Document, image_paths: list[Path]) -> None:
	"""Додає зображення до розділу завдання"""
	if not image_paths:
		return

	doc.add_heading('ІЛЮСТРАЦІЇ', 2)
	for img_path in image_paths:
		try:
			doc.add_paragraph(f"Файл: {img_path.name}")
			doc.add_picture(str(img_path), width=Inches(6))
			last_paragraph = doc.paragraphs[-1]
			last_paragraph.paragraph_format.space_after = Pt(12)
		except Exception as e:
			doc.add_paragraph(f"Не вдалося додати зображення {img_path.name}: {e}")


def create_report() -> Path:
	"""Створює звіт у форматі DOCX"""

	ensure_packages()

	py_files = get_python_files()
	task_results = {}

	print("\nЗбір результатів виконання...")

	for task_file in py_files:
		input_text = TEST_INPUTS.get(task_file.name)
		output = capture_output(task_file, input_text=input_text)

		if output is None:
			task_results[task_file.name] = "(Без виводу)"
			print(f"⊘ {task_file.name}: результат не отримано")
		else:
			task_results[task_file.name] = output if output.strip() else "(Без виводу)"
			print(f"✓ {task_file.name}: OK")

	# Створюємо документ
	doc = Document()

	# Налаштування стилів
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = Pt(12)

	# Титульна сторінка
	title = doc.add_heading('ЗВІТ', 0)
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER

	subtitle = doc.add_paragraph('з лабораторної роботи №5')
	subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
	subtitle.runs[0].font.size = Pt(14)
	subtitle.runs[0].font.bold = True

	theme_p = doc.add_paragraph()
	theme_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	theme_run = theme_p.add_run(
		'Тема: Побудова рекомендаційної системи відеопереглядів'
	)
	theme_run.font.size = Pt(12)
	theme_run.font.italic = True

	doc.add_paragraph()

	subject = doc.add_paragraph('Дисципліна: Машинне навчання, обробка великих масивів даних')
	subject.alignment = WD_ALIGN_PARAGRAPH.CENTER

	doc.add_paragraph()
	doc.add_paragraph()

	info = doc.add_paragraph()
	info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	info_run = info.add_run(
		f'Виконав: Косарєв Валерій Геннадійович\n'
		f'Напрямок: 051 Економіка\n'
		f'Дата: {datetime.now().strftime("%d.%m.%Y")}'
	)
	info_run.font.size = Pt(12)

	doc.add_page_break()

	# Зміст
	toc = doc.add_heading('ЗМІСТ', 1)
	toc.alignment = WD_ALIGN_PARAGRAPH.CENTER

	for i, file_path in enumerate(py_files, 1):
		filename = file_path.name
		if filename in TASK_DESCRIPTIONS:
			task_info = TASK_DESCRIPTIONS[filename]
			doc.add_paragraph(f"{i}. {task_info['title']}", style='List Bullet')

	doc.add_page_break()

	# Теоретичні відомості
	doc.add_heading('1. ТЕОРЕТИЧНІ ВІДОМОСТІ', 1)
	theory_text = (
		"Тема. Побудова рекомендаційної системи відеопереглядів.\n\n"
		"Мета. Набути умінь та навичок розробки рекомендаційних систем для порталу відопереглядів за допомогою метрики та використання мови програмування Python.\n\n"
		"Рекомендаційні системи використовуються для пропозиції користувачам контенту, "
		"який може їм сподобатися на основі їхньої поведінки та переваг. Основні підходи включають:\n\n"
		"1. Методи, що базуються на вмісті (content-based filtering)\n"
		"2. Методи, що базуються на спільнотах (collaborative filtering)\n"
		"3. Гібридні методи\n\n"
		"Ключові компоненти включають: завантаження даних, нормалізацію, обчислення метрик подібності та генерацію рекомендацій."
	)
	doc.add_paragraph(theory_text)
	doc.add_page_break()

	# Основна частина
	doc.add_heading('2. РЕЗУЛЬТАТИ ВИКОНАННЯ ЗАВДАНЬ', 1)

	for task_num, file_path in enumerate(py_files, 1):
		filename = file_path.name

		if filename in TASK_DESCRIPTIONS:
			task_info = TASK_DESCRIPTIONS[filename]
			doc.add_heading(f"{task_num}. {task_info['title']}", 1)
		else:
			doc.add_heading(f"{task_num}. {filename}", 1)
			task_info = {'description': 'Виконання завдання', 'topic': 'Python'}

		# 1. Короткий огляд
		doc.add_heading('1. Короткий огляд', 2)
		overview = f"{task_info['description']}\nТема: {task_info['topic']}"
		overview_para = doc.add_paragraph(overview)
		overview_para.paragraph_format.left_indent = Inches(0.5)
		doc.add_paragraph()

		# 2. Вихідний код
		doc.add_heading('2. Вихідний код:', 2)
		code_content = read_file_content(file_path)
		add_formatted_code(doc, code_content)
		doc.add_paragraph()

		# 3. Результат виконання
		doc.add_heading('3. Результат виконання:', 2)
		output = task_results.get(filename, "(Без виводу)")
		for line in output.split('\n'):
			if line.strip():
				p = doc.add_paragraph(line)
				p.paragraph_format.left_indent = Inches(0.5)
				if p.runs:
					p.runs[0].font.name = 'Consolas'
					p.runs[0].font.size = Pt(9)
					p.runs[0].font.color.rgb = RGBColor(0, 100, 0)
			else:
				doc.add_paragraph()

		doc.add_paragraph()

		# Додавання зображень, якщо вони відповідають завданню
		task_images = find_task_images(filename)
		if task_images:
			add_task_images_section(doc, task_images)
			doc.add_paragraph()

		# 4. Висновки
		doc.add_heading('4. Висновки:', 2)
		conclusions = get_task_conclusions(filename)
		conclusion_para = doc.add_paragraph(conclusions)
		conclusion_para.paragraph_format.left_indent = Inches(0.5)

		if task_num < len(py_files):
			doc.add_page_break()

	# Висновки
	doc.add_page_break()
	doc.add_heading('3. ВИСНОВКИ', 1)
	general_conclusion = (
		"Під час виконання лабораторної роботи №5 було сформовано практичні навички "
		"розробки рекомендаційних систем для порталу відеопереглядів. Засвоєно методи "
		"завантаження та нормалізації даних, розраховані метрики подібності користувачів "
		"(евклідова, Манхеттенська, Max-метрика) та реалізована система генерації рекомендацій "
		"на основі обраної метрики. Отримані результати демонструють практичне застосування "
		"алгоритмів рекомендаційних систем для реальних завдань порталів відеопереглядів."
	)
	doc.add_paragraph(general_conclusion)

	# Контрольні запитання
	doc.add_heading('4. ВІДПОВІДІ НА КОНТРОЛЬНІ ЗАПИТАННЯ', 1)

	doc.add_heading('4.1 Поняття рекомендаційної системи.', 2)
	q1 = (
		"Рекомендаційна система — це програмна система, яка пропонує користувачам елементи контенту, "
		"що можуть їм сподобатися, на основі їхньої поведінки та переваг. Основні типи рекомендаційних систем: "
		"1) Фільтрування на основі змісту (content-based) — рекомендуються елементи, подібні до тих, які користувач уже використовував; "
		"2) Спільнотне фільтрування (collaborative filtering) — рекомендуються елементи, якими користуються користувачі з подібними смаками; "
		"3) Гібридні підходи — комбінація декількох методів. Рекомендаційні системи широко використовуються у порталах відеопереглядів, "
		"соціальних мережах, інтернет-магазинах та інших платформах для підвищення задоволення користувачів і збільшення залучення контенту."
	)
	doc.add_paragraph(q1)

	doc.add_heading('4.2 Алгоритми нормалізації даних.', 2)
	q2 = (
		"Нормалізація даних — це процес приведення числових даних до спільної шкали для порівняння. "
		"Основні методи нормалізації: "
		"1) Min-Max нормалізація (масштабування): кожне значення x трансформується за формулою (x - min) / (max - min), "
		"що приводить значення до діапазону [0, 1]; "
		"2) Z-score нормалізація (стандартизація): (x - середнє) / стандартне_відхилення, що дає розподіл зі середнім 0 та дисперсією 1; "
		"3) Десяткова нормалізація: ділення на відповідний ступінь 10; "
		"4) Вектор нормалізація (одиничне векторне масштабування): ділення кожного значення на норму вектора. "
		"Вибір методу залежить від природи даних і завдання машинного навчання."
	)
	doc.add_paragraph(q2)

	doc.add_heading('4.3 Види метрик (Евклідова, метрика Манхеттен, Max-метрика).', 2)
	q3 = (
		"Метрики обчислюють відстань між двома точками в просторі ознак, що дозволяє оцінити подібність користувачів: "
		"1) Евклідова метрика (L2-норма): sqrt(sum((xi - yi)²)) — обчислює прямолінійну відстань, найчастіше використовується у машинному навчанні; "
		"2) Манхеттенська метрика (L1-норма, міська відстань): sum(|xi - yi|) — обчислює суму абсолютних різниць координат, менш чутлива до викидів; "
		"3) Max-метрика (Чебишева, L∞-норма): max(|xi - yi|) — використовує максимальну різницю координат, зручна для рівномірного розподілу. "
		"Євклідова метрика найпоширеніша у рекомендаційних системах, але вибір метрики залежить від специфіки даних та вимог завдання."
	)
	doc.add_paragraph(q3)

	output_path = CURRENT_DIR / f'Звіт_Лабораторна_5_{datetime.now().strftime("%Y%m%d")}.docx'
	doc.save(str(output_path))

	print(f"\n✓ Звіт успішно створено: {output_path}")
	print(f"✓ Проаналізовано файлів: {len(py_files)}")
	print(f"✓ Розмір файлу: {output_path.stat().st_size / 1024:.2f} KB")

	return output_path


if __name__ == '__main__':
	print("=" * 70)
	print("Генерація звіту з лабораторної роботи №5")
	print("=" * 70)

	try:
		create_report()
		print()
		print("=" * 70)
		print("Звіт готовий!")
		print("=" * 70)
	except Exception as e:
		print(f"✗ Помилка при створенні звіту: {e}")
		import traceback
		traceback.print_exc()
